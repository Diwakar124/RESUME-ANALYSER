import streamlit as st
import pickle
import docx
import PyPDF2
import re
import io
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

# -------------------- CONFIG --------------------
st.set_page_config(page_title="AI Resume Analyzer Pro", page_icon="🚀", layout="wide")

# -------------------- CUSTOM CSS --------------------
st.markdown("""
<style>
.main {
    background-color: #0e1117;
    color: white;
}
h1, h2, h3 {
    color: #00f2ff;
    text-align: center;
}
.stButton>button {
    background-color: #00f2ff;
    color: black;
    border-radius: 10px;
    padding: 10px;
    font-weight: bold;
}
</style>
""", unsafe_allow_html=True)

# -------------------- HEADER --------------------
st.markdown("""
<h1>🚀 AI Resume Analyzer</h1>
<p style='text-align:center; font-size:18px;'>
Analyze your resume, get ATS score & match with job descriptions instantly.
</p>
""", unsafe_allow_html=True)

# -------------------- FEATURE CARDS --------------------
col1, col2, col3 = st.columns(3)

with col1:
    st.info("📄 Resume Classification")

with col2:
    st.info("📊 ATS Score")

with col3:
    st.info("💼 Job Matching")

st.markdown("---")

# -------------------- LOAD MODELS --------------------
@st.cache_resource
def load_models():
    svc_model = pickle.load(open('clf.pkl', 'rb'))
    tfidf = pickle.load(open('tfidf.pkl', 'rb'))
    le = pickle.load(open('encoder.pkl', 'rb'))
    return svc_model, tfidf, le

svc_model, tfidf, le = load_models()

# -------------------- FUNCTIONS --------------------
def clean_resume(txt):
    txt = re.sub('http\\S+\\s', ' ', txt)
    txt = re.sub('RT|cc', ' ', txt)
    txt = re.sub('#\\S+\\s', ' ', txt)
    txt = re.sub('@\\S+', ' ', txt)
    txt = re.sub('[%s]' % re.escape("!\"#$%&'()*+,-./:;<=>?@[\\]^_`{|}~"), ' ', txt)
    txt = re.sub(r'[^\\x00-\\x7f]', ' ', txt)
    txt = re.sub('\\s+', ' ', txt)
    return txt

def extract_text(file):
    ext = file.name.split('.')[-1].lower()
    text = ""

    if ext == 'pdf':
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text += page.extract_text() or ""

    elif ext == 'docx':
        doc = docx.Document(file)
        for para in doc.paragraphs:
            text += para.text + '\n'

    elif ext == 'txt':
        try:
            text = file.read().decode('utf-8')
        except:
            text = file.read().decode('latin-1')

    return text

def predict_category(text):
    cleaned = clean_resume(text)
    vec = tfidf.transform([cleaned]).toarray()
    pred = svc_model.predict(vec)
    return le.inverse_transform(pred)[0]

def extract_skills(text):
    skills_db = ["python","java","c++","sql","machine learning","deep learning",
                 "data analysis","pandas","numpy","tensorflow","keras",
                 "excel","power bi","tableau","aws","html","css","javascript"]
    text = text.lower()
    return list(set([s for s in skills_db if s in text]))

def ats_score(text):
    keywords = ["experience","project","skills","education","internship","certification","achievement"]
    score = sum(1 for k in keywords if k in text.lower())
    return int((score/len(keywords))*100)

def recommend_role(skills):
    if "machine learning" in skills:
        return "Machine Learning Engineer"
    elif "python" in skills and "sql" in skills:
        return "Data Analyst"
    elif "html" in skills and "css" in skills:
        return "Frontend Developer"
    return "General Software Role"

def match_resume_job(resume, job_desc):
    tfidf_local = TfidfVectorizer(stop_words='english')
    vectors = tfidf_local.fit_transform([resume, job_desc])
    score = cosine_similarity(vectors[0:1], vectors[1:2])[0][0]
    return round(score * 100, 2)

def missing_keywords(resume, job_desc):
    return list(set(job_desc.lower().split()) - set(resume.lower().split()))[:20]

# -------------------- INPUT --------------------
tab1, tab2 = st.tabs(["📂 Upload Resume", "📌 Job Description"])

with tab1:
    uploaded_file = st.file_uploader("Upload Resume", type=["pdf","docx","txt"])

with tab2:
    job_desc = st.text_area("Paste Job Description", height=200)

# -------------------- OUTPUT --------------------
if uploaded_file:
    text = extract_text(uploaded_file)

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("📄 Resume Content")
        st.text_area("", text, height=400)

    with col2:
        st.subheader("📊 Analysis Results")

        st.success(f"🎯 Category: {predict_category(text)}")

        skills = extract_skills(text)
        st.write("🛠 Skills:", skills)

        score = ats_score(text)
        st.progress(score)
        st.write(f"ATS Score: {score}%")

        st.success(f"💼 Role: {recommend_role(skills)}")

        if job_desc:
            match_score = match_resume_job(text, job_desc)
            st.progress(int(match_score))
            st.write(f"Match Score: {match_score}%")

            st.write("⚠️ Missing Keywords:", missing_keywords(text, job_desc))


pdf = generate_pdf(
    category,
    skills,
    score,
    role,
    match_score if job_desc else None
)

st.download_button(
    label="📄 Download Report as PDF",
    data=pdf,
    file_name="resume_analysis.pdf",
    mime="application/pdf"
)
def improve_resume(resume_text, job_desc):
    resume_words = set(resume_text.lower().split())
    job_words = set(job_desc.lower().split())

    missing = job_words - resume_words

    # pick important words (limit)
    missing = list(missing)[:15]

    suggestions = f"""
    🔧 Resume Improvement Suggestions:

    - Add these important keywords:
      {', '.join(missing)}

    - Include relevant projects related to job description
    - Highlight experience using these keywords
    - Add measurable achievements (e.g., improved accuracy by 20%)

    ✨ Suggested Resume Line:
    "Worked on projects involving {', '.join(missing[:5])} to align with job requirements."
    """

    return suggestions
if job_desc:
    match_score = match_resume_job(text, job_desc)

    st.progress(int(match_score))
    st.write(f"Match Score: {match_score}%")

    missing = missing_keywords(text, job_desc)
    st.write("⚠️ Missing Keywords:", missing if missing else "None 🎉")

    st.info(suggestions(match_score))

    if match_score < 60:
        st.subheader("🚀 Improve Your Resume")

    improved_text = improve_resume(text, job_desc)

    st.warning(improved_text)

    new_resume = text + "\n\n" + improved_text

    st.subheader("📄 Improved Resume Draft")
    st.text_area("", new_resume, height=300)



# -------------------- CONFIG --------------------
st.set_page_config(page_title="AI Resume Analyzer Pro", page_icon="🚀", layout="wide")

# -------------------- CSS --------------------
st.markdown("""
<style>
.main {background-color:#0e1117;color:white;}
h1,h2,h3 {color:#00f2ff;text-align:center;}
.stButton>button {background:#00f2ff;color:black;border-radius:10px;}
</style>
""", unsafe_allow_html=True)

# -------------------- HEADER --------------------
st.markdown("""
<h1>🚀 AI Resume Analyzer</h1>
<p style='text-align:center;'>Analyze resume + ATS + Job Match</p>
""", unsafe_allow_html=True)

# -------------------- LOAD MODEL --------------------
@st.cache_resource
def load_models():
    return (
        pickle.load(open('clf.pkl','rb')),
        pickle.load(open('tfidf.pkl','rb')),
        pickle.load(open('encoder.pkl','rb'))
    )

svc_model, tfidf, le = load_models()

# -------------------- FUNCTIONS --------------------
def clean_resume(txt):
    txt = re.sub('http\\S+\\s',' ',txt)
    txt = re.sub('\\s+',' ',txt)
    return txt

def extract_text(file):
    ext = file.name.split('.')[-1]
    text = ""
    if ext == 'pdf':
        reader = PyPDF2.PdfReader(file)
        for p in reader.pages:
            text += p.extract_text() or ""
    elif ext == 'docx':
        doc = docx.Document(file)
        for para in doc.paragraphs:
            text += para.text
    elif ext == 'txt':
        text = file.read().decode('utf-8', errors='ignore')
    return text

def predict_category(text):
    vec = tfidf.transform([clean_resume(text)]).toarray()
    return le.inverse_transform(svc_model.predict(vec))[0]

def extract_skills(text):
    skills = ["python","sql","machine learning","html","css","java"]
    return [s for s in skills if s in text.lower()]

def ats_score(text):
    keys = ["experience","project","skills","education"]
    return int(sum(k in text.lower() for k in keys)/len(keys)*100)

def recommend_role(skills):
    if "machine learning" in skills:
        return "ML Engineer"
    if "python" in skills:
        return "Data Analyst"
    return "Software Role"

def match_resume_job(r,j):
    vec = TfidfVectorizer(stop_words='english').fit_transform([r,j])
    return round(cosine_similarity(vec[0:1],vec[1:2])[0][0]*100,2)

def missing_keywords(r,j):
    return list(set(j.lower().split()) - set(r.lower().split()))[:10]

def suggestions(score):
    return "Good Match ✅" if score>70 else "Improve Resume ⚠️"

def improve_resume(r,j):
    miss = missing_keywords(r,j)
    return f"Add keywords: {', '.join(miss)}"

# -------------------- PDF --------------------
def generate_pdf(cat,skills,score,role,match):
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer)
    style = getSampleStyleSheet()
    content = [
        Paragraph(f"Category: {cat}",style["Normal"]),
        Paragraph(f"Skills: {skills}",style["Normal"]),
        Paragraph(f"ATS: {score}%",style["Normal"]),
        Paragraph(f"Role: {role}",style["Normal"]),
    ]
    if match:
        content.append(Paragraph(f"Match: {match}%",style["Normal"]))
    doc.build(content)
    buffer.seek(0)
    return buffer

# -------------------- INPUT --------------------
tab1,tab2 = st.tabs(["Upload Resume","Job Description"])

with tab1:
    uploaded_file = st.file_uploader("Upload Resume")

with tab2:
    job_desc = st.text_area("Paste JD")

# -------------------- OUTPUT --------------------
if uploaded_file:
    text = extract_text(uploaded_file)

    col1,col2 = st.columns(2)

    with col1:
        st.text_area("Resume",text,height=300)

    with col2:
        category = predict_category(text)
        skills = extract_skills(text)
        score = ats_score(text)
        role = recommend_role(skills)

        st.success(f"Category: {category}")
        st.write("Skills:",skills)
        st.progress(score)
        st.write(f"ATS: {score}%")
        st.success(f"Role: {role}")

        match_score = None

        if job_desc:
            match_score = match_resume_job(text,job_desc)
            st.progress(int(match_score))
            st.write(f"Match: {match_score}%")
            st.write("Missing:",missing_keywords(text,job_desc))
            st.info(suggestions(match_score))

            if match_score < 60:
                st.warning(improve_resume(text,job_desc))
                st.text_area("Improved Resume",
                             text+"\n\n"+improve_resume(text,job_desc),
                             height=200)

        # PDF Download
        pdf = generate_pdf(category,skills,score,role,match_score)
        st.download_button("📄 Download PDF",pdf,"report.pdf")

# -------------------- FOOTER --------------------
st.markdown("---")
st.markdown("<center>Made by Diwakar 🚀</center>",unsafe_allow_html=True)